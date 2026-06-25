"""
Resume parser — handles PDF and DOCX files.
Returns raw text plus detected section headings.
"""
from __future__ import annotations

import io
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Common resume section headings we try to detect
SECTION_PATTERNS = [
    r"summary|objective|profile",
    r"experience|employment|work history",
    r"education|academic",
    r"skills|technical skills|core competencies",
    r"projects|portfolio",
    r"certifications?|licenses?",
    r"awards?|honors?|achievements?",
    r"publications?|research",
    r"volunteer|community",
    r"languages?",
    r"interests?|hobbies",
    r"references?",
]
SECTION_RE = re.compile(
    r"(?im)^(" + "|".join(SECTION_PATTERNS) + r")\s*[:\-]?\s*$"
)


def _detect_sections(text: str) -> list[str]:
    found: list[str] = []
    for match in SECTION_RE.finditer(text):
        label = match.group(1).strip().title()
        if label not in found:
            found.append(label)
    return found


# ── PDF ───────────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> tuple[str, list[str]]:
    """Extract text from a PDF file. Returns (raw_text, detected_sections)."""
    try:
        import pdfplumber  # lazy import — only needed if PDF uploaded

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=3)
                if text:
                    pages.append(text)

        raw = "\n".join(pages)
        return raw, _detect_sections(raw)

    except ImportError:
        raise RuntimeError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        )
    except Exception as exc:
        logger.exception("PDF parsing failed")
        raise ValueError(f"Could not parse PDF: {exc}") from exc


# ── DOCX ─────────────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> tuple[str, list[str]]:
    """Extract text from a DOCX file. Returns (raw_text, detected_sections)."""
    try:
        from docx import Document  # python-docx

        doc = Document(io.BytesIO(file_bytes))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also pull text from tables (skills tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        raw = "\n".join(paragraphs)
        return raw, _detect_sections(raw)

    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    except Exception as exc:
        logger.exception("DOCX parsing failed")
        raise ValueError(f"Could not parse DOCX: {exc}") from exc


# ── Public entry point ────────────────────────────────────────────────────────

def parse_resume(filename: str, file_bytes: bytes) -> tuple[str, list[str]]:
    """
    Route to the correct parser based on file extension.
    Returns (raw_text, detected_sections).
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_bytes)
    elif suffix in (".docx", ".doc"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Please upload a PDF or DOCX."
        )
