"""
Export service — generates ATS-compliant PDF and DOCX from optimized text.
Intentionally keeps formatting minimal so ATS parsers never choke on it.
"""
from __future__ import annotations

import io
import re
import logging

logger = logging.getLogger(__name__)

# ── Helpers ───────────────────────────────────────────────────────────────────

def _is_section_heading(line: str) -> bool:
    """Heuristic: ALL CAPS short line, or line ending with a colon."""
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.endswith(":") and len(stripped) < 50:
        return True
    if stripped.isupper() and len(stripped.split()) <= 5:
        return True
    return False


def _parse_lines(text: str) -> list[str]:
    return [l.rstrip() for l in text.splitlines()]


# ── PDF export ────────────────────────────────────────────────────────────────

def generate_pdf(optimized_text: str, target_role: str) -> bytes:
    """
    Build an ATS-friendly single-column PDF using ReportLab.
    Uses only standard Helvetica fonts — no custom fonts needed.
    """
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        heading_style = ParagraphStyle(
            "SectionHeading",
            fontName="Helvetica-Bold",
            fontSize=11,
            spaceAfter=2,
            textColor=colors.HexColor("#1a1a1a"),
            borderPaddingBottom=2,
        )
        body_style = ParagraphStyle(
            "Body",
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1a1a1a"),
            alignment=TA_LEFT,
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=body_style,
            leftIndent=14,
            bulletIndent=0,
        )

        story = []
        lines = _parse_lines(optimized_text)

        for line in lines:
            if not line.strip():
                story.append(Spacer(1, 4))
                continue

            if _is_section_heading(line):
                story.append(Spacer(1, 6))
                story.append(Paragraph(line.strip().upper(), heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#555555"), spaceAfter=4))
            elif line.strip().startswith(("•", "-", "*", "·")):
                clean = re.sub(r"^[•\-\*·]\s*", "", line.strip())
                story.append(Paragraph(f"• {clean}", bullet_style))
            else:
                story.append(Paragraph(line.strip(), body_style))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")
    except Exception as exc:
        logger.exception("PDF generation failed")
        raise RuntimeError(f"PDF generation error: {exc}") from exc


# ── DOCX export ───────────────────────────────────────────────────────────────

def generate_docx(optimized_text: str, target_role: str) -> bytes:
    """
    Build an ATS-friendly DOCX using python-docx.
    Single-column, no text boxes, no tables — pure paragraphs.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Remove default section margins — set tight ATS margins
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Remove default blank paragraph
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)

        lines = _parse_lines(optimized_text)

        for line in lines:
            if not line.strip():
                doc.add_paragraph("")
                continue

            if _is_section_heading(line):
                p = doc.add_paragraph()
                run = p.add_run(line.strip().upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                # Underline divider via bottom border on paragraph
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "555555")
                pBdr.append(bottom)
                pPr.append(pBdr)

            elif line.strip().startswith(("•", "-", "*", "·")):
                clean = re.sub(r"^[•\-\*·]\s*", "", line.strip())
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(clean)
                run.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.font.size = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as exc:
        logger.exception("DOCX generation failed")
        raise RuntimeError(f"DOCX generation error: {exc}") from exc
